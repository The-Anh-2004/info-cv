# với cv ứng viên tải lên, input là link, chỉ đọc tên, ngành nghề, số điện thoại, email, địa chỉ, ngày sinh, giới tính.
from ultralytics import YOLO
import torch
import cv2
import os
from get_address import extract_address
from extract_address import extract_city, extract_district, extract_ward
import easyocr
import fitz
import re
import docx
# from underthesea import ner
import spacy
### OCR ###
import torch
import cv2
import os
import matplotlib.pyplot as plt
# import shutil
from OCR_server.test_craft import *
from OCR_server.craft import *
from OCR_server.pipeline import *
from OCR_server.inference import *
from OCR_server.load_model import *
# import glob
# import numpy as np
# from pdf2image import convert_from_path
# from flask import Flask,request
# import json
# from PIL import Image
# import uuid
# import io
# from vietocr.tool.predictor import Predictor
# from vietocr.tool.config import Cfg
# import argparse
# import torch.nn as nn
# import torch.backends.cudnn as cudnn
# from torch.autograd import Variable
# import urllib
from unidecode import unidecode
# import multiprocessing

def chuyen_cau_khong_dau_chu_thuong(cau):
    # Chuyển đổi văn bản có dấu thành văn bản không dấu
    cau_khong_dau = unidecode(cau)
    
    # Chuyển câu thành chữ thường
    cau_khong_dau_chu_thuong = cau_khong_dau.lower()
    
    return cau_khong_dau_chu_thuong
def process_cropped_image(cropped_img):
    list_crop_line = crop_image_line(cropped_img,craft,args,refine_net)
    return list_crop_line
    
def extractText(image): 
    list_crop_line = crop_image_line_info(image,craft,args,refine_net)

    _output, all_info_box = recog(list_crop_line,ocr_model_1)
    text = _output["title"] + ' ' +  _output["text"]
    return text



det_box_model, ocr_model, craft, refine_net, args, ocr_model_1 = get_model()


### OCR ###

map_label = {0: 'avatar',
             1: 'block',
             2: 'infor',
             3: 'job_title',
             4: 'name'}
 
def extract_text(image, det_model):
    image = cv2.cvtColor(image, cv2.COLOR_BGR2GRAY)
    image = cv2.cvtColor(image, cv2.COLOR_GRAY2RGB)
    output_path = 'results'
    det_model = det_model
    boxes_list, label_list, detect_image, confs = det_model(image, output_path= output_path, return_result=True)
    name = []
    infor = []
    for i in range(len(boxes_list)):
        if label_list[i] == torch.tensor(4.) :
            cropped_img = image[int(boxes_list[i][1]):int(boxes_list[i][3]), int(boxes_list[i][0]):int(boxes_list[i][2]), :]
            cv2.imwrite(os.path.join('folder_data/' + str(i) + '.jpg'), cropped_img)
            text = extractText(cropped_img)
            name.append(text)
        if label_list[i] == torch.tensor(3.) :
            cropped_img = image[int(boxes_list[i][1]):int(boxes_list[i][3]), int(boxes_list[i][0]):int(boxes_list[i][2]), :]
            cv2.imwrite(os.path.join('./folder_data/' + str(i) + '.jpg'), cropped_img)
            text = extractText(cropped_img)
            infor.append(text)    
        if label_list[i] == torch.tensor(2.):
            cropped_img = image[int(boxes_list[i][1]):int(boxes_list[i][3]), int(boxes_list[i][0]):int(boxes_list[i][2]), :]
            cv2.imwrite(os.path.join('folder_data/' + str(i) + '.jpg'), cropped_img)
            text = extractText(cropped_img)
            infor.append(text)
        if label_list[i] == torch.tensor(1.) :
            cropped_img = image[int(boxes_list[i][1]):int(boxes_list[i][3]), int(boxes_list[i][0]):int(boxes_list[i][2]), :]
            cv2.imwrite(os.path.join('folder_data/' + str(i) + '.jpg'), cropped_img)
            text = extractText(cropped_img)
            infor.append(text)    

    return name, infor
def extract_job_title(image, det_model):
    image = cv2.cvtColor(image, cv2.COLOR_BGR2GRAY)
    image = cv2.cvtColor(image, cv2.COLOR_GRAY2RGB)
    output_path = './results'
    det_model = det_model
    boxes_list, label_list, detect_image_job, confs = det_model(image, output_path= output_path, return_result=True)
    cv2.imwrite(os.path.join('folder_data/detect_image.jpg'), detect_image_job)
    job_title = []
    for i in range(len(boxes_list)):
        if label_list[i] == torch.tensor(3.) :
            cropped_img = image[int(boxes_list[i][1]):int(boxes_list[i][3]), int(boxes_list[i][0]):int(boxes_list[i][2]), :]
            cv2.imwrite(os.path.join('./folder_data/' + str(i) + '.jpg'), cropped_img)
            text = extractText(cropped_img)
            job_title.append(text)

    return job_title

def pdf2text(path):
    doc = fitz.open(path)
    text = ""
    for page in doc:
        textpage = page.get_textpage()      
        text += textpage.extractTEXT()     
    return text

def doc2text(path):
    doc = docx.Document(path)
    '''
    except:
        if doc is None:
            convert_to_docx(filename)
            file = filename.replace('.doc', '.docx')
            os.remove(filename)
            doc = docx.Document(file)
    '''
    fullText = []
    for para in doc.paragraphs:
        fullText.append(para.text)
    infor =  '\n'.join(fullText)
    return infor



def check_text(text1, text2):
    count = 0
    for t1 in text1:
        if t1 in text2:
            count = count + 1
    rate = 2*count/(len(text1)+len(text2))
    return rate


def extract_gender(text):
    text = text.replace('Việt Nam', '')
    text = text.replace('việt nam', '')
    text = text.replace('Vietnam', '')
    text = text.replace('vietnam', '')
    patterns = ['Nam', 'Nữ', 'Male', 'Female', 'nam', 'nữ', 'NAM', 'NỮ']
    count = 0
    for pattern in patterns:
        if pattern in text:
            gender = pattern
            count = count + 1
            break
    if count == 0:
        gender = 'Khác'
    return gender


def extract_email(text):
    patterns = [
        r'[\w\.-]+@[\w\.-]+\.\w+',
        r'[\w\.-]+\s*@\s*[\w\.-]+\s*\.\s*\w+'
    ]
    text = text.replace(' ogmail', '@gmail')
    text = text.replace('qgmail', '@gmail')
    text = text.replace('0gmail', '@gmail')
    text = text.replace('0gmail.com', '@gmail.com')
    text = text.replace(' gmail.com', '@gmail.com')
    text = text.replace('gqmail.com', 'gmail.com')
    text = text.replace('gmial.com', 'gmail.com')
    text = text.replace('gmaik.com', 'gmail.com')
    text = text.replace('gmajl.com', 'gmail.com')
    text = text.replace('gmal.com', 'gmail.com')
    text = text.replace('gmai.com', 'gmail.com')
    text = text.replace('gmailcom', 'gmail.com')
    text = text.replace('gmai1', 'gmail')
    text = text.replace('gmai|', 'gmail')
    text = text.replace('gmaiI', 'gmail')
    text = text.replace('gmaii', 'gmail')
    text = text.replace('gnail.com', 'gmail.com')
    text = text.replace('gmall.com', 'gmail.com')
    text = text.replace('agmail', 'gmail')
    text = text.replace('agmail.com', 'gmail.com')
    text = text.replace('Agmail.com', 'gmail.com')
    text = text.replace('A gmail.com', '@gmail.com')
    text = text.replace('@gmaii.com', '@gmail.com')
    text = text.replace('@gmai1.com', '@gmail.com')
    text = text.replace('@gmai|.com', '@gmail.com')
    text = text.replace('@gmaiI.com', '@gmail.com')
    text = text.replace('@gmaIl.com', '@gmail.com')  # L viết hoa
    text = text.replace('@gmaLl.com', '@gmail.com')  # LL
    text = text.replace('gnail.com', 'gmail.com')
    text = text.replace('gmai1.com', 'gmail.com')
    text = text.replace('gmai|.com', 'gmail.com')
    text = text.replace('gmaiI.com', 'gmail.com')
    text = text.replace('gmaIl.com', 'gmail.com')
    text = text.replace('gmaLl.com', 'gmail.com')
    text = text.replace('gmaiil.com', 'gmail.com')
    text = text.replace('gmaiil.com', 'gmail.com')
    text = text.replace('gmai.com', 'gmail.com')
    text = text.replace('gmali.com', 'gmail.com')
    text = text.replace('gmal.com', 'gmail.com')
    text = text.replace('gamil.com', 'gmail.com')
    text = text.replace('gmaul.com', 'gmail.com')
    text = text.replace('gmai|.com', 'gmail.com')
    text = text.replace('gmaill.com', 'gmail.com')
    text = text.replace('gmaii.com', 'gmail.com')
    text = text.replace('gmaik.com', 'gmail.com')
    text = text.replace('gnmai.com', 'gmail.com')
    text = text.replace('gnail.com', 'gmail.com')
    text = text.replace('qmail.com', 'gmail.com')
    text = text.replace('gmajl.com', 'gmail.com')
    text = text.replace('gmalil.com', 'gmail.com')
    text = text.replace('gmall.com', 'gmail.com')
    text = text.replace('gmali.com', 'gmail.com')
    text = text.replace('gma1l.com', 'gmail.com')
    text = text.replace('gmauil.com', 'gmail.com')
    text = text.replace('gmai|.com', 'gmail.com')
    text = text.replace('gma1l.com', 'gmail.com')
    text = text.replace('gmail..com', 'gmail.com')
    text = text.replace('gmaill.com', 'gmail.com')
    text = text.replace('@gmail.com.com', '@gmail.com')
    text = text.replace('@gmai1.com', '@gmail.com')
    text = text.replace('gmail,com', 'gmail.com')
    text = text.replace('gmail.corn', 'gmail.com')
    text = text.replace('gmail.cpm', 'gmail.com')
    text = text.replace('gmail.con', 'gmail.com')
    text = text.replace('gmail.co', 'gmail.com')
    text = text.replace('gmail.coim', 'gmail.com')
    text = text.replace('gmail.cim', 'gmail.com')
    text = text.replace('gmaii.c0m', 'gmail.com')
    text = text.replace('gmail.cm', 'gmail.com')
    text = text.replace('gmall.com', 'gmail.com')
    text = text.replace('gmaiI.com', 'gmail.com')
    text = text.replace('gmai1.com', 'gmail.com')
    text = text.replace('gmaii.com', 'gmail.com')
    text = text.replace('gmaik.com', 'gmail.com')
    text = text.replace('gmaiil.com', 'gmail.com')
    text = text.replace('gmal.com', 'gmail.com')
    text = text.replace('gmail.c0m', 'gmail.com')
    text = text.replace('gmail.col', 'gmail.com')
    text = text.replace('gnail.com', 'gmail.com')
    text = text.replace('gma1l.com', 'gmail.com')
    text = text.replace('gmail.vom', 'gmail.com')
    text = text.replace('gmail.xom', 'gmail.com')
    text = text.replace('gmail,com', 'gmail.com')
    text = text.replace('gmail.cpm', 'gmail.com')
    text = text.replace('gamil.com', 'gmail.com')
    text = text.replace('gmaii.com', 'gmail.com')
    text = text.replace('gmiail.com', 'gmail.com')
    text = text.replace('gmail.com.vn', 'gmail.com')
    text = text.replace('gmaill.com', 'gmail.com')
    text = text.replace('gnail.com', 'gmail.com')
    text = re.sub(r'\b(\w+)\s+gmail\.com\b', r'\1@gmail.com', text)
    for pattern in patterns:
        if re.findall(pattern, text):
            email = re.findall(pattern, text)
            return email[0]
    return None

def extract_phone(text):
    patterns = ['[0-9]{10}', '[0-9]{5} [0-9]{5}',
                '[0-9]{4} [0-9]{3} [0-9]{3}',
                '[0-9]{4}.[0-9]{3}.[0-9]{3}',
                '[0-9]{3} [0-9]{3} [0-9]{4}',
                '[0-9]{3}.[0-9]{3}.[0-9]{4}',
                '[0-9]{4}-[0-9]{3}-[0-9]{3}',
                '[0-9]{3}-[0-9]{3}-[0-9]{4}',
                r'(\(\+84\)[\s\-\.]?\d{3}[\s\-\.]?\d{3}[\s\-\.]?\d{3})',
                r'(\(\+84\)[\s\-\.]?\d{9,10})',
                r'(\+84[\s\-\.]?\d{9,10})',
                r'(84[\s\-\.]?\d{9,10})']
    for pattern in patterns:
        match = re.findall(pattern, text)
        if match:
            phone = match[0]
            # Loại bỏ các ký tự thừa: (, ), khoảng trắng, ., -
            phone_clean = re.sub(r'[\s\.\-\(\)]', '', phone)
            # Chuẩn hóa đầu số
            if phone_clean.startswith('+84'):
                phone_clean = '0' + phone_clean[3:]
            elif phone_clean.startswith('84'):
                phone_clean = '0' + phone_clean[2:]
            # Chỉ lấy 10 số cuối nếu dư (trường hợp nhập thừa)
            if len(phone_clean) > 10:
                phone_clean = phone_clean[-10:]
            return phone_clean
    return None
    

def extract_date_of_birth(text):
    text = text.replace('年 ', '年')
    text = text.replace(' 年', '年')
    text = text.replace('月 ', '月')
    text = text.replace(' 年', '年')
    text = text.replace('日 ', '日')
    text = text.replace(' 日', '日')
    patterns = [r'[A-Za-z]+\s\d{1,2},\s\d{4}', r'[A-Za-z]+\s\d{1,2}(?:st|nd|rd|th)?\s\d{4}', r'\s\d{1,2} [A-Za-z]+ \s\d{4}',
                r'\d{1,2}\s(?:January|February|March|April|May|June|July|August|September|October|November|December)\s\d{4}', 
                r'[0-9]{4}年[0-9]{2}月[0-9]{2}日', r'[A-Za-z]+\s\d{1,2}(?:St|Nd|Rd|Th)?\s\d{4}', r'\d{2}\s[/.-]\s\d{2}\s[/.-]\s\d{4}',
                r'\d{1}[/]\d{2}[/]\d{4}', r'\d{1}[/]\d{1}[/]\d{4}', r'\d{1}[-]\d{2}[-]\d{4}', r'\d{1}[.]\d{1}[.]\d{4}', r'\d{1}[-]\d{1}[-]\d{4}',
                r'\d{2}[.]\d{1}[.]\d{4}', r'\d{2}[-]\d{1}[-]\d{4}',  r'\d{1}[.]\d{2}[.]\d{4}', r'\d{2}[-]\d{2}[-]\d{4}',  r'\d{4}[-]\d{2}[-]\d{2}',
                r'\d{2}[/]\d{2}[/]\d{4}', r'\d{4}[/]\d{2}[/]\d{2}', r'\d{4}[.]\d{2}[.]\d{2}', r'\d{2}[/]\d{1}[/]\d{4}',  r'\d{2}[.]\d{2}[.]\d{4}']
    birth = ''
    age = ''
    for pattern in patterns:
        if re.findall(pattern, text):
            for date in re.findall(pattern, text):
                if re.findall('[0-9]{4}', date) and int(re.findall('[0-9]{4}', date)[0]) < 2006:
                    birth = date
                    age = 2023 - int(re.findall('[0-9]{4}', birth)[0])
                    break
    return birth, age

def show_text(arr):
    if len(arr) > 0:
        text = arr[0]
        for i in range(1, len(arr)):
            text = text+', '+arr[i]
    else:
        text = ''
    return text

def extract_infor_img(
    name, info_list,
    cat_id, cat_name,
    user_id,
    job_title,
    city_id=None, city_name=None,
    district_id=None, district_name=None,
    ward_id=None, ward_name=None
):
    def safe_str(x):
        return str(x) if x not in [None, 'None'] else ''

    infors = {
        'email': '',
        'phone': '',
        'birthday': '',
        'gender': '',
        'age': '',
        'name': name[0] if name else '',
        'title_cv': job_title or '',
        'cat_id': safe_str(cat_id),
        'cat_name': cat_name or '',
        'address': '',
        'city_id': safe_str(city_id),
        'city_name': city_name or '',
        'district_id': safe_str(district_id),
        'district_name': district_name or '',
        'ward_id': safe_str(ward_id),
        'ward_name': ward_name or '',
        'user_id': user_id or ''
    }

    # Lấy email, phone
    for text in info_list:
        if not infors['email']:
            email_found = extract_email(text)
            if email_found:
                infors['email'] = email_found
        if not infors['phone']:
            phone_found = extract_phone(text)
            if phone_found:
                infors['phone'] = phone_found
        if infors['email'] and infors['phone']:
            break
    # Ngày sinh & tuổi
    for text in info_list:
        birth, age = extract_date_of_birth(text)
        if birth:
            infors['birthday'] = birth
            infors['age'] = safe_str(age)
            break
    # Giới tính
    for text in info_list:
        gender = extract_gender(text)
        if gender and gender.lower() not in ['khác', 'other']:
            infors['gender'] = gender
            break
    # Địa chỉ
    for text in info_list:
        address = extract_address(text)
        if address:
            address = address.replace(' /', '/').replace('/ ', '/').replace(' / ', '/').replace('. ', '.')
            infors['address'] = address
            break
    return infors


def extract_infor(text, user_id):
    phone = ''
    email = ''
    birthday = ''
    gender = ''
    title_cv = ''
    fullname = ''
    address = ''
    age = ''
    infors = {}
    if extract_email(text) != None:
        email = extract_email(text)
    if extract_phone(text) != None:
        phone = extract_phone(text)
    gender = extract_gender(text)
    birthday, age = extract_date_of_birth(text)
    nlp_ner = spacy.load("model")
    doc = nlp_ner(text)
    for entity in doc.ents:
        if entity.label_ == 'PERSON':
            fullname = entity.text
            break
    address = extract_address(text)
    (city_id, city_name,text_no_city) = extract_city(text, "api-base365.City.json")
    (district_id, district_name,text_no_district) = extract_district(text_no_city, city_id, "api-base365.District.json")
    ward_id, ward_name = extract_ward(text_no_district, city_id, "api-base365.Ward.json")
    infors['email'] = email
    infors['phone'] = phone
    infors['birthday'] = birthday
    infors['gender'] = gender
    infors['title_cv'] = title_cv
    infors['fullname'] = fullname
    infors['name'] = fullname
    infors['age'] = age
    infors['address'] = address
    infors['city_id'] = city_id
    infors['city_name'] = city_name 
    infors['district_id'] = district_id
    infors['district_name'] = district_name
    infors['ward_id'] = ward_id
    infors['ward_name'] = ward_name
    infors['user_id'] = user_id
    return infors


#def extract_infor_from_cv(path, user_id):
